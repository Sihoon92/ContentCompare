"""Word Raw Extractor — python-docx 로 .docx 를 physical_raw 로 변환.

문단/표를 **문서 등장 순서대로** 뽑는다. 해석은 하지 않는다 — 어느 문단이
heading 이고 어느 표가 규격표인지는 LLM 이 raw json 을 보고 판단한다.

순서 보존이 중요한 이유
----------------------
python-docx 의 ``document.paragraphs`` 와 ``document.tables`` 는 서로 분리된
목록이라 그대로 쓰면 "문단 다음에 표" 같은 본문 흐름이 사라진다. 그래서 본문
XML(``document.element.body``)의 자식 순서를 직접 훑어 paragraph(``w:p``)와
table(``w:tbl``)을 등장 순서대로 처리한다.

python-docx 는 선택적 의존성이므로 import 를 함수 안으로 지연한다.
"""

from __future__ import annotations

import logging
import os
from typing import Any, Optional

from .models import RawWordBlock, RawWordDocument

logger = logging.getLogger(__name__)


def extract_word_raw(path: str) -> RawWordDocument:
    """docx 파일 경로 → :class:`RawWordDocument` (문서 순서 보존)."""
    try:
        import docx  # python-docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:  # pragma: no cover - 환경 의존
        raise RuntimeError(
            "python-docx 가 필요합니다. pip install python-docx"
        ) from exc

    file_name = os.path.basename(path)
    logger.info("[RawWord] 열기: %s", os.path.abspath(path))
    document = docx.Document(path)

    doc = RawWordDocument(file_name=file_name)
    order = 0
    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]  # 네임스페이스 제거
        if tag == "p":
            para = Paragraph(child, document)
            text = (para.text or "").strip()
            if not text:
                continue  # 빈 문단(레이아웃용)은 건너뛴다
            order += 1
            doc.blocks.append(
                RawWordBlock(
                    block_id=f"w_b{order:03d}",
                    order=order,
                    type="paragraph",
                    text=text,
                    style=_paragraph_style(para),
                )
            )
        elif tag == "tbl":
            table = Table(child, document)
            rows = _table_rows(table)
            if not rows:
                continue
            order += 1
            doc.blocks.append(
                RawWordBlock(
                    block_id=f"w_b{order:03d}",
                    order=order,
                    type="table",
                    rows=rows,
                )
            )

    logger.info("[RawWord] 완료: 블록 %d개", len(doc.blocks))
    return doc


def _paragraph_style(para) -> Optional[dict[str, Any]]:
    """문단 스타일 정보. 알 수 있는 값만 담고, 전부 없으면 None."""
    info: dict[str, Any] = {}

    style = getattr(para, "style", None)
    if style is not None and getattr(style, "name", None):
        info["style_name"] = style.name

    # 굵게/글자크기: 첫 run 기준(문단 대표값). run 이 없으면 생략.
    bold, size = _first_run_format(para)
    if bold is not None:
        info["bold"] = bold
    if size is not None:
        info["font_size"] = size

    return info or None


def _first_run_format(para):
    """문단의 첫 비어있지 않은 run 의 (bold, font_size_pt) 추정."""
    for run in para.runs:
        if not (run.text or "").strip():
            continue
        bold = run.bold
        size = None
        if run.font is not None and run.font.size is not None:
            # EMU/Pt 객체 → 포인트 수치.
            size = float(run.font.size.pt)
        return bold, size
    return None, None


def _table_rows(table) -> list[list[str]]:
    """표 → 셀 텍스트 2D 리스트(공백 정돈)."""
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([" ".join((cell.text or "").split()) for cell in row.cells])
    return rows
