"""samples/spec_en.docx 생성 + 기준 엑셀 값 채우기 + 정답지 생성.

``samples/자표준_규격서.docx`` 의 영어판이다. 원래는 단락 11 + 표 6행짜리 소형 픽스처였고
그 부분은 ``LEGACY_*`` 로 **그대로 보존**한다(기존 골든셋 라벨이 이 문장들에 걸려 있다).

여기에 기준 엑셀 27~106행(규격 항목 80건)에 대응하는 영문 내용을 얹어, **평가 가능한**
교차언어 비교 픽스처로 확장했다. 세 산출물이 **하나의 ``CASES`` 테이블**에서 함께 나온다:

1. ``samples/자표준문서.xlsx`` 27~106행의 값 칸(F~N)  — ``fill_reference()``
2. ``samples/spec_en.docx``                          — ``build_word()``
3. ``golden/spec_en_골든셋.jsonl`` (104건 정답지)     — ``write_golden()``

문서만 바꾸고 정답을 안 고치는 드리프트가 구조적으로 불가능하게 하려는 것이다
(``scripts/make_synthetic_targets.py`` 와 같은 설계, ``golden/README.md`` 참고).

**3~26행은 건드리지 않는다.** ``golden/자표준_골든셋.jsonl`` 27건이 그 행의 값과
"기준 문서는 단위 열(K)이 비어 있다"는 라벨 전제에 걸려 있다. 그 행들은 이미 문서에 있는
``LEGACY_*`` 문장과 대조해 정답지에만 등재한다(``place="pre"``).

표기 관습 차이는 의도적이다 — 언어 탓으로 돌릴 수 있는 차이를 남겨야 교차언어 검증이 된다:

- 온도 단위를 ``℃``(U+2103) 대신 ``°C``(도 기호 + C)로 쓴다.
- 습도를 ``33~53%RH`` 대신 ``33 to 53 %RH`` 로 쓴다.
- 전류를 ``1495mA`` 대신 ``1.495A`` 로 쓴다.

실행:

    pip install python-docx openpyxl
    python scripts/make_en_sample.py                 # 셋 다
    python scripts/make_en_sample.py --skip-reference # 엑셀은 건드리지 않음
"""

from __future__ import annotations

import argparse
import json
import os
import sys

from docx import Document

OUT_DOCX = "samples/spec_en.docx"
OUT_GOLDEN = "golden/spec_en_골든셋.jsonl"
REF_XLSX = "samples/자표준문서.xlsx"
REF_SHEET = "데이터"
REF_FIRST_NEW_ROW = 27      # 신규 규격 항목이 시작하는 행 (3~26 은 불가침)

# 엑셀 열 인덱스
COL = {"seq": 2, "class": 3, "mid": 4, "sub": 5, "lower": 6, "target": 7, "upper": 8,
       "spec": 9, "qual": 10, "unit": 11, "level": 12, "remark": 13, "method": 14}


# ---------------------------------------------------------------------------
# 기존 픽스처 (건드리지 않는다 — 골든셋 27건이 이 문장들을 참조한다)
# ---------------------------------------------------------------------------

LEGACY_PARAGRAPHS = [
    ("h1", "Battery Cell Specification (Summary)"),
    ("p", "This document summarizes the specification items of the master standard "
          "document in narrative form."),
    ("h2", "1. Electrical Specifications"),
    ("p", "This specification conforms to Battery Approval Specification ver 4.7 "
          "(SEC Req. ver.4.7)."),
    ("p", "The rated charging voltage is 4.55V."),
    ("p", "The discharge cut-off voltage shall be 3.0V."),
    ("p", "The nominal voltage is 3.85V."),
    ("p", "The maximum charging current shall not exceed 1.495A."),
    ("p", "The evaluation ambient humidity shall be within the range of 33 to 53 %RH, "
          "with a center value of 43 %RH."),
]

# c27932e 가 손으로 넣은 4구간 문단. 별도 단락이 아니라 **한 문단 안의 <w:br/>** 다 —
# 그 형태 자체가 F5 1:N 판정의 회귀 케이스이므로 재생성해도 반드시 같은 모양이어야 한다.
LEGACY_CHARGE_TEMP = [
    "Charge temperature ranges:",
    "-5~5℃, 0.1C(4.55V)",
    "5~12℃, 0.3C(4.55V)",
    "12~15℃, 0.7C(4.55V)",
    "15~45℃, 1.2C(4.20V)",
]

LEGACY_TABLE_HEADING = "2. Key Specification Table"
LEGACY_TABLE = [
    ("Item", "Specification", "Unit"),
    ("Nominal capacity", "1150", "mAh"),
    ("Rated capacity", "1150", "mAh"),
    ("Standard charging current", "230", "mA"),
    ("Maximum discharging current", "1150", "mA"),
    ("Standard ambient temperature", "21 ~ 29 (center 25)", "°C"),
]

FOOTER = "* Items not specified in this document shall be agreed upon separately."

# 실제 규격서에서 옮겨온 "2.3 General Specification" 표. 셀 안의 ``\n`` 은 <w:br/> 로 들어간다.
#
# ⚠️ 이 표는 앞 절(1. Electrical Specifications)의 서술과 **일부러 어긋난다** — 같은 항목을
# 문서 두 곳이 다르게 말하는 상황이 실문서에 흔하고, 그때 파이프라인이 무엇을 하는지가
# 정확히 F5 1:N 판정의 시험대이기 때문이다. 어긋나는 곳은 네 군데다:
#   공칭전압      3.85V (1절)      vs 3.89V (2.3.3)
#   표준용량 Typ. 1180mAh (6절)    vs 1,185mAh (2.3.1)
#   12~15 구간    0.7C (충전온도 문단) vs 0.8C (2.3.23)
#   15~45 구간    1.2C(4.20V)      vs 1.3C/1.1C/0.8C (2.3.23)
# 기준에 대응 행이 없는 항목(E/D 697WH/L, cutoff 23mA, 45~55 구간)도 일부러 남겼다 —
# 기준에 없는 것을 있다고 우기지 않는지 보는 쪽 시험이다.
GENERAL_SPEC_HEADING = "2.3 General Specification"
GENERAL_SPEC_TABLE = [
    ("No.", "Item", "Definition", "Specifications"),
    ("2.3.1", "Capacity", "Rated\nStd. discharge capacity and criterion of C-rate",
     "1,150mAh"),
    ("", "", "Typ.", "1,185mAh"),
    ("2.3.2", "Energy Density (E/D)",
     "(Min capacity X Nominal Voltage) / max Volume X1000\n(Without side type)\n"
     "* Standard E/D :1st Vendor\n* Max Volume at SOC for ex-factory condition.",
     "697WH/L"),
    ("2.3.3", "Nominal voltage",
     "Mean operation voltage during standard discharge after standard charge", "3.89V"),
    ("2.3.4", "Charging Method", "CC-CV(constant voltage with limited current)", "CC-CV"),
    ("2.3.5", "Standard charge", "Charge voltage\nCharge current\nCutoff condition",
     "4.55V\n230mA (0.2C)\n23mA (0.02C)"),
    ("2.3.6", "Standard discharge", "Discharge current\nCutoff voltage",
     "230mA(0.2C)\n3.0V"),
    ("2.3.23", "Operating Protocol in different Temperatures",
     "Cell charge protocol by different temperature range",
     "-5~5°C, 0.1C(4.55V)\n5~12°C, 0.3C(4.55V)\n12~15°C, 0.8C(4.55V)\n"
     "15~45°C, 1.3C(4.10V)\n1.1C(4.28V)\n0.8C(4.55V)\n45~55°C, 0.8C(4.28V)"),
]


# ---------------------------------------------------------------------------
# 마스터 테이블
# ---------------------------------------------------------------------------

def C(**kw):
    """케이스 하나. 빠진 키는 기본값으로 채운다."""
    base = dict(seq=None, row=None, ko="", en="", lower=None, target=None, upper=None,
                unit=None, qual=None, level="C", remark=None, method=None,
                place=None, cells=None, text=None,
                expected="missing", mismatch=(), reason="")
    base.update(kw)
    return base


# place 값 = 워드 문서에서 어디에 실리는가
#   pre    : 이미 LEGACY_* 에 있음 (엑셀 3~26행 — 값 칸을 건드리지 않는 행)
#   method : 3. Measurement and Sampling — 서술 문단
#   dim    : 4. Dimensions and Weight  (Item/Min/Typ/Max/Unit 5열 표)
#   bullet : 5. Appearance and Workmanship 불릿
#   elec   : 6. Electrical Performance (Item/Specification/Unit/Method 4열 표)
#   rate   : 7. Rate and Temperature Discharge (Condition/Retention 2열 표)
#   env    : 8. Environmental Test Results (Test item/Criteria/Unit 3열 표)
#   para   : 9. Storage Characteristics — 서술 문단
#   store  : 9. Storage Characteristics (Condition/Residual/Recovered 3열 표, 항목명 열 없음)
#   ilt    : 10. Mass Production Inspection — "Term: value" 정의형 문단
#   note   : 11. Notes 불릿
#   None   : 문서에 없음 (missing)

CASES = [
    # ---- 기존 3~26행: 값 칸 불가침. 이미 문서에 있는 문장과 대조만 한다 ----------
    C(row=3, ko="고객 표준 버전", en="Approval specification version",
      qual="배터리승인규격 ver 4.7\nSEC Req. ver.4.7", place="pre",
      text="This specification conforms to Battery Approval Specification ver 4.7 (SEC Req. ver.4.7).",
      expected="match", reason="승인규격 버전 문구가 기준과 동일"),
    C(row=4, ko="공칭용량", en="Nominal capacity", lower=1150, place="pre",
      text="Nominal capacity | 1150 | mAh",
      expected="match", reason="1150 일치"),
    C(row=5, ko="정격용량", en="Rated capacity", lower=1150, place="pre",
      text="Rated capacity | 1150 | mAh",
      expected="match", reason="1150 일치"),
    C(row=6, ko="정격충전전압", en="Rated charging voltage", target=4.55, place="pre",
      text="The rated charging voltage is 4.55V.",
      expected="match", reason="4.55 일치"),
    C(row=7, ko="공칭전압", en="Nominal voltage", target=3.89, place="pre",
      text="The nominal voltage is 3.85V.",
      expected="mismatch", mismatch=("target_value",),
      reason="대상 문서가 자기모순 — 1절은 3.85V, 2.3.3 은 3.89V. 기준 3.89 와 어긋나는 진술이 있으므로 mismatch"),
    C(row=8, ko="deltaOCV", en="Delta OCV", expected="missing",
      reason="영문 문서에 △OCV 항목 없음"),
    C(row=9, ko="충전방법", en="Charging method", place="pre",
      text="2.3.4 | Charging Method | CC-CV(constant voltage with limited current) | CC-CV",
      expected="unknown",
      reason="대상 2.3.4 에 CC-CV 가 있으나 기준 9행은 값 칸이 비어 있어(3~26행 불가침) 대조 불가"),
    C(row=10, ko="충전방법(SOC28->SOC64)", en="Charging method (SOC28->SOC64)",
      expected="missing", reason="영문 문서에 해당 구간 충전 방법 없음"),
    C(row=11, ko="표준충전전류", en="Standard charging current", target=230, place="pre",
      text="Standard charging current | 230 | mA",
      expected="match", reason="230 일치"),
    C(row=12, ko="최대충전전류", en="Maximum charging current", target=1495, place="pre",
      text="The maximum charging current shall not exceed 1.495A.",
      expected="unknown",
      reason="기준 1495 는 단위 열이 비어 있어 mA 인지 확정 불가 — 1.495A 와 등가인지 판단보류"),
    C(row=13, ko="최대방전전류", en="Maximum discharging current", target=1150, place="pre",
      text="Maximum discharging current | 1150 | mA",
      expected="match", reason="1150 일치"),
    C(row=14, ko="방전종지전압", en="Discharge cut-off voltage", target=3, place="pre",
      text="The discharge cut-off voltage shall be 3.0V.",
      expected="match", reason="기준 3 과 표기만 다름(3 vs 3.0)"),
    C(row=15, ko="표준환경온도", en="Standard ambient temperature",
      lower=21, target=25, upper=29, place="pre",
      text="Standard ambient temperature | 21 ~ 29 (center 25) | °C",
      expected="match", reason="21/25/29 일치, 단위 표기만 ℃ vs °C"),
    C(row=16, ko="평가환경습도", en="Evaluation ambient humidity",
      lower=33, target=43, upper=53, place="pre",
      text="The evaluation ambient humidity shall be within the range of 33 to 53 %RH, "
           "with a center value of 43 %RH.",
      expected="match", reason="33/43/53 일치"),
    C(row=17, ko="충전환경온도", en="Charging ambient temperature",
      lower=-5, target=35, upper=85, expected="missing",
      reason="영문 문서에 충전환경온도 단일 항목 없음 (구간표만 있음)"),
    C(row=18, ko="방전환경온도", en="Discharging ambient temperature",
      lower=-30, target=30, upper=90, expected="missing",
      reason="영문 문서에 없음"),
    C(row=19, ko="평가환경온도", en="Evaluation ambient temperature",
      lower=21, target=24, upper=28, expected="missing", reason="영문 문서에 없음"),
    C(row=20, ko="1개월저장온도", en="Storage temperature (1 month)",
      lower=-10, target=35, upper=80, expected="missing", reason="영문 문서에 없음"),
    C(row=21, ko="3개월저장온도", en="Storage temperature (3 months)",
      lower=-10, target=35, upper=70, expected="missing", reason="영문 문서에 없음"),
    C(row=22, ko="1년저장온도", en="Storage temperature (1 year)",
      lower=-10, target=35, upper=55, expected="missing", reason="영문 문서에 없음"),
    C(row=23, ko="충전온도범위(-5~5℃)", en="Charge temperature range (-5 to 5 °C)",
      lower=-5, upper=5, qual="충전전류 0.1C, 충전전압 4.55V", place="pre",
      text="-5~5℃, 0.1C(4.55V)",
      expected="match", reason="구간·전류·전압 모두 일치"),
    C(row=24, ko="충전온도범위(5~12℃)", en="Charge temperature range (5 to 12 °C)",
      lower=5, upper=12, qual="충전전류 0.3C, 충전전압 4.55V", place="pre",
      text="5~12℃, 0.3C(4.55V)",
      expected="match", reason="구간·전류·전압 모두 일치"),
    C(row=25, ko="충전온도범위(12~15℃)", en="Charge temperature range (12 to 15 °C)",
      lower=12, upper=15, qual="충전전류 0.7C, 충전전압 4.55V", place="pre",
      text="12~15°C, 0.8C(4.55V)",
      expected="mismatch", mismatch=("qualitative_spec",),
      reason="대상 문서가 자기모순 — 충전온도 문단은 0.7C, 2.3.23 은 0.8C. 기준 0.7C 와 어긋나는 진술이 있음"),
    C(row=26, ko="충전온도범위(15~45℃)", en="Charge temperature range (15 to 45 °C)",
      lower=15, upper=45, qual="충전전류 1.2C, 충전전압 4.20V", place="pre",
      text="15~45°C, 1.3C(4.10V)",
      expected="mismatch", mismatch=("qualitative_spec",),
      reason="대상 문서가 자기모순 — 충전온도 문단은 1.2C(4.20V), 2.3.23 은 같은 구간에 1.3C(4.10V)·1.1C(4.28V)·0.8C(4.55V) 세 조건"),

    # ---- 신규 27~106행: 값 칸을 여기서 채운다 ----------------------------------
    # 기본사양 --------------------------------------------------------------
    C(seq=90, ko="두께 측정방법", en="Thickness measurement method",
      qual="평판 가압법, 하중 300gf, 유지 10초", level="B",
      remark="측정 지그 MS-JIG-02 사용", method="MS-TH-001",
      place="method",
      text="Cell thickness shall be measured by the flat plate method under a load of "
           "300 gf, held for 10 seconds.",
      expected="match", reason="측정법·하중·유지시간 모두 일치"),
    C(seq=220, ko="Recording 조건", en="Recording condition",
      qual="1초 간격, 16bit 분해능", level="C",
      remark="설비 로그 보존 3년", method="MS-DAQ-004",
      expected="missing", reason="영문 문서에 계측 기록 조건 서술 없음"),
    C(seq=230, ko="누액 검사", en="Leakage inspection",
      qual="60℃ 24시간 방치 후 전해액 누액 없을 것", level="A",
      remark="전수 검사", method="MS-LK-002",
      place="bullet",
      text="No electrolyte leakage shall be observed after 24 hours at 60 °C.",
      expected="match", reason="조건(60℃·24h)과 판정 기준 일치"),
    C(seq=240, ko="수명평가 시료선정", en="Sample selection for cycle life",
      qual="Lot 당 5셀 무작위 추출", level="C",
      remark="Lot 정의는 별도 합의", method="MS-SP-001",
      place="method",
      text="Samples for the life evaluation shall be selected in accordance with the "
           "internal sampling plan.",
      expected="unknown", reason="대상은 '내부 샘플링 계획'만 언급하고 수량(5셀/Lot)을 밝히지 않아 대조 불가"),

    # 일반특성 --------------------------------------------------------------
    C(seq=250, ko="외관", en="Appearance",
      qual="찍힘·스크래치·오염 없을 것", level="B",
      remark="육안 검사, 조도 500lx 이상", method="MS-AP-001",
      place="bullet",
      text="The cell surface shall be free from dents, scratches and contamination.",
      expected="match", reason="세 결함 항목이 기준과 동일"),
    C(seq=260, ko="만충전두께 3", en="Fully charged thickness",
      lower=4.35, target=4.45, upper=4.55, unit="mm", level="A",
      remark="충전 직후 1시간 이내 측정", method="MS-TH-001",
      place="dim", cells=("4.35", "4.45", "4.55", "mm"),
      expected="match", reason="하한·중심·상한 모두 일치"),
    C(seq=270, ko="출하충전두께1", en="Shipping charged thickness",
      lower=4.20, target=4.30, upper=4.40, unit="mm", level="A",
      remark="출하 SOC 30% 기준", method="MS-TH-001",
      place="dim", cells=("4.20", "4.30", "4.35", "mm"),
      expected="mismatch", mismatch=("upper_limit",),
      reason="기준 상한 4.40 vs 대상 4.35"),
    C(seq=280, ko="폭치수", en="Width", lower=62.80, target=63.00, upper=63.20,
      unit="mm", level="B", remark="Tab 제외", method="MS-DIM-001",
      place="dim", cells=("62.80", "63.00", "63.20", "mm"),
      expected="match", reason="치수 3개 모두 일치"),
    C(seq=290, ko="총고치수", en="Overall height", lower=90.30, target=90.60, upper=90.90,
      unit="mm", level="B", remark="Strip Tape 제외", method="MS-DIM-001",
      place="dim", cells=("90.30", "90.60", "90.90", "mm"),
      expected="match", reason="치수 3개 모두 일치"),
    C(seq=300, ko="총고치수2 (Strip Tape까지의 총길이)", en="Overall height incl. strip tape",
      lower=91.00, target=91.40, upper=91.80, unit="mm", level="C",
      remark="Strip Tape 포함", method="MS-DIM-001",
      expected="missing", reason="영문 문서에 Strip Tape 포함 치수 없음"),
    C(seq=310, ko="셀방총고2", en="Cell body height", lower=89.50, target=89.80,
      upper=90.10, unit="mm", level="C", remark="파우치 실링부 제외", method="MS-DIM-001",
      expected="missing", reason="영문 문서에 셀 본체 높이 없음"),
    C(seq=320, ko="Cell 좌/우 총고 편차", en="Left/right height deviation",
      upper=0.15, unit="mm", level="B", remark="좌우 3점 평균 차", method="MS-DIM-002",
      expected="missing", reason="영문 문서에 좌우 편차 항목 없음"),
    C(seq=330, ko="중량", en="Weight", lower=21.5, target=22.0, upper=22.5,
      unit="g", level="B", remark="포장재 제외", method="MS-WT-001",
      place="dim", cells=("21.5", "22.0", "22.5", "g"),
      expected="match", reason="중량 3개 모두 일치"),
    C(seq=340, ko="출하충전 IR", en="Shipping internal resistance",
      lower=15, target=30, upper=45, unit="mΩ", level="B",
      remark="1kHz AC 측정", method="MS-IR-001",
      place="elec", cells=("15 ~ 45 (typ. 30)", "mΩ", "MS-IR-001"),
      expected="match", reason="하한 15 / 중심 30 / 상한 45 일치"),
    C(seq=350, ko="출하충전 OCV", en="Shipping open circuit voltage",
      lower=3.80, target=3.83, upper=3.86, unit="V", level="B",
      remark="출하 후 24시간 경과 기준", method="MS-OCV-001",
      place="elec", cells=("3.83", "-", "MS-OCV-001"),
      expected="unknown", reason="대상이 중심치 3.83 만 제시하고 상·하한을 밝히지 않아 범위 대조 불가"),
    C(seq=360, ko="출하충전율", en="Shipping state of charge",
      lower=28, target=30, upper=32, unit="%", level="A",
      remark="항공 운송 규정 연계", method="MS-SOC-001",
      place="elec", cells=("33 ~ 37 (typ. 35)", "%", "MS-SOC-001"),
      expected="mismatch", mismatch=("lower_limit", "target_value", "upper_limit"),
      reason="기준 28/30/32 vs 대상 33/35/37"),
    C(seq=370, ko="PP stripHeight", en="PP strip height",
      lower=2.8, target=3.0, upper=3.2, unit="mm", level="C",
      remark="Tape 부착 후 측정", method="MS-DIM-003",
      place="bullet",
      text="The PP strip height is approximately 3.",
      expected="unknown", reason="대상이 단위를 밝히지 않고 'approximately' 로만 서술해 mm 대조 불가"),
    C(seq=380, ko="양극 Tab 위치", en="Cathode tab position",
      lower=12.4, target=12.7, upper=13.0, unit="mm", level="B",
      remark="셀 좌측 기준", method="MS-DIM-004",
      place="dim", cells=("12.4", "12.7", "13.0", "mm"),
      expected="match", reason="Tab 위치 3개 모두 일치"),
    C(seq=390, ko="음극 Tab 위치", en="Anode tab position",
      lower=24.9, target=25.2, upper=25.5, unit="mm", level="B",
      remark="셀 좌측 기준", method="MS-DIM-004",
      place="dim", cells=("24.9", "25.0", "25.5", "mm"),
      expected="mismatch", mismatch=("target_value",),
      reason="기준 중심 25.2 vs 대상 25.0"),
    C(seq=400, ko="바인더 접착력 1", en="Binder adhesion strength",
      lower=12, target=18, unit="gf/mm", level="B",
      remark="90도 박리 시험", method="MS-AD-001",
      expected="missing", reason="영문 문서에 바인더 접착력 항목 없음"),
    C(seq=410, ko="실링 강도 (상부)", en="Sealing strength (top)",
      lower=45, target=55, unit="N/15mm", level="A",
      remark="180도 인장", method="MS-SL-001",
      place="bullet",
      text="The top sealing strength shall be at least 45 N/15mm (typically 55 N/15mm).",
      expected="match", reason="하한 45 / 대표 55 일치"),
    C(seq=420, ko="실링 강도 (하부)", en="Sealing strength (bottom)",
      lower=40, target=50, unit="N/15mm", level="A",
      remark="180도 인장", method="MS-SL-001",
      expected="missing", reason="영문 문서에 하부 실링 강도 없음"),
    C(seq=430, ko="Sepa 결착력", en="Separator adhesion",
      lower=8, target=12, unit="gf/mm", level="B",
      remark="상온 24시간 후 측정", method="MS-AD-002",
      place="bullet",
      text="The separator adhesion shall be no less than 0.08 N/mm.",
      expected="unknown", reason="대상 단위 N/mm 와 기준 gf/mm 의 환산 근거가 문서에 없어 등가 판단 불가"),

    # 신뢰성특성 / 환경특성 -------------------------------------------------
    C(seq=440, ko="온습도1 후 △V(2hr)", en="Voltage drop after damp heat (2 h)",
      upper=0.020, unit="V", level="A", remark="60℃/90%RH 48시간", method="MS-EV-101",
      place="env", cells=("<= 0.020", "V"),
      expected="match", reason="상한 0.020 일치"),
    C(seq=450, ko="온습도1 후 △V(24hr)", en="Voltage drop after damp heat (24 h)",
      upper=0.030, unit="V", level="A", remark="60℃/90%RH 48시간", method="MS-EV-101",
      place="env", cells=("<= 0.030", "V"),
      expected="match", reason="상한 0.030 일치"),
    C(seq=460, ko="온습도1 후 △R", en="Resistance rise after damp heat",
      upper=15, unit="%", level="B", remark="초기 IR 대비", method="MS-EV-101",
      expected="missing", reason="영문 문서에 온습도 후 저항 상승률 없음"),
    C(seq=470, ko="온습도1 후 △T(2hr두께) 3", en="Thickness increase after damp heat (2 h)",
      upper=5, unit="%", level="A", remark="초기 두께 대비", method="MS-EV-101",
      place="env", cells=("<= 6", "%"),
      expected="mismatch", mismatch=("upper_limit",),
      reason="기준 상한 5% vs 대상 6%"),
    C(seq=480, ko="온습도1 후 외관", en="Appearance after damp heat",
      qual="부풀음·누액·변색 없을 것", level="A", remark="육안 검사", method="MS-EV-101",
      place="env", cells=("No swelling, leakage or discoloration", "-"),
      expected="match", reason="세 결함 항목이 기준과 동일"),
    C(seq=490, ko="온습도1 후 용량(0.2C)", en="Capacity after damp heat (0.2C)",
      lower=95, unit="%", level="B", remark="초기 용량 대비", method="MS-EV-101",
      place="env", cells=("No significant capacity loss", "-"),
      expected="unknown", reason="대상이 정성 서술만 하고 수치(95%)를 제시하지 않아 대조 불가"),
    C(seq=500, ko="온습도1 후 용량(1.0C)", en="Capacity after damp heat (1.0C)",
      lower=93, unit="%", level="B", remark="초기 용량 대비", method="MS-EV-101",
      place="env", cells=(">= 93", "%"),
      expected="match", reason="하한 93% 일치"),
    C(seq=510, ko="온습도1 후 7일방치OCV", en="OCV after damp heat and 7 day rest",
      lower=3.75, unit="V", level="A", remark="상온 7일 방치", method="MS-EV-101",
      place="env", cells=(">= 3.75", "V"),
      expected="match", reason="하한 3.75V 일치"),
    C(seq=520, ko="온습도1 후 한계평가", en="Limit test after damp heat",
      qual="발화·폭발 없을 것", level="A", remark="안전 항목", method="MS-EV-901",
      place="bullet",
      text="No fire and no explosion shall occur in the limit test after the damp heat test.",
      expected="match", reason="발화·폭발 없음 기준 일치"),
    C(seq=530, ko="온습도 후 해체", en="Teardown after damp heat",
      qual="음극 표면 리튬 석출 없을 것", level="A", remark="해체 분석", method="MS-EV-902",
      expected="missing", reason="영문 문서에 해체 분석 결과 없음"),
    C(seq=540, ko="열충격1 후 △V  (2hr)", en="Voltage drop after thermal shock (2 h)",
      upper=0.025, unit="V", level="A", remark="-30℃↔85℃ 50사이클", method="MS-EV-201",
      place="env", cells=("<= 0.025", "V"),
      expected="match", reason="상한 0.025 일치"),
    C(seq=550, ko="열충격1 후 △V  (24hr)", en="Voltage drop after thermal shock (24 h)",
      upper=0.035, unit="V", level="A", remark="-30℃↔85℃ 50사이클", method="MS-EV-201",
      place="env", cells=("<= 0.030", "V"),
      expected="mismatch", mismatch=("upper_limit",),
      reason="기준 상한 0.035V vs 대상 0.030V"),
    C(seq=560, ko="열충격1 후 △R", en="Resistance rise after thermal shock",
      upper=20, unit="%", level="B", remark="초기 IR 대비", method="MS-EV-201",
      place="env", cells=("<= 20", "%"),
      expected="match", reason="상한 20% 일치"),
    C(seq=570, ko="열충격1 후 △T  (직후두께) 3", en="Thickness increase after thermal shock (immediate)",
      upper=8, unit="%", level="A", remark="시험 직후 측정", method="MS-EV-201",
      place="env", cells=("<= 8", "%"),
      expected="match", reason="상한 8% 일치"),
    C(seq=580, ko="열충격1 후 △T  (2hr두께) 3", en="Thickness increase after thermal shock (2 h)",
      upper=6, unit="%", level="A", remark="2시간 후 측정", method="MS-EV-201",
      expected="missing", reason="영문 문서에 2시간 후 두께 증가율 없음"),
    C(seq=590, ko="열충격1 후 외관", en="Appearance after thermal shock",
      qual="부풀음·누액·변색 없을 것", level="A", remark="육안 검사", method="MS-EV-201",
      place="env", cells=("No swelling, leakage or discoloration", "-"),
      expected="match", reason="세 결함 항목이 기준과 동일"),
    C(seq=600, ko="열충격1 후 용량(0.2C)", en="Capacity after thermal shock (0.2C)",
      lower=93, unit="%", level="B", remark="초기 용량 대비", method="MS-EV-201",
      place="env", cells=(">= 93", "%"),
      expected="match", reason="하한 93% 일치"),
    C(seq=610, ko="열충격1 후 용량(1.0C)", en="Capacity after thermal shock (1.0C)",
      lower=91, unit="%", level="B", remark="초기 용량 대비", method="MS-EV-201",
      expected="missing", reason="영문 문서에 1.0C 용량 유지율 없음"),
    C(seq=620, ko="열충격1 후 7일방치  OCV", en="OCV after thermal shock and 7 day rest",
      lower=3.72, unit="V", level="A", remark="상온 7일 방치", method="MS-EV-201",
      place="env", cells=(">= 3.72", "V"),
      expected="match", reason="하한 3.72V 일치"),
    C(seq=630, ko="열충격1 후 한계평가", en="Limit test after thermal shock",
      qual="발화·폭발 없을 것", level="A", remark="안전 항목", method="MS-EV-901",
      place="bullet",
      text="No fire and no explosion shall occur in the limit test after the thermal shock test.",
      expected="match", reason="발화·폭발 없음 기준 일치"),
    C(seq=640, ko="열충격 후 해체", en="Teardown after thermal shock",
      qual="음극 표면 리튬 석출 없을 것", level="A", remark="해체 분석", method="MS-EV-902",
      expected="missing", reason="영문 문서에 해체 분석 결과 없음"),
    C(seq=650, ko="고온동작△R (2hr)", en="Resistance rise in high temperature operation (2 h)",
      upper=18, unit="%", level="B", remark="45℃ 동작", method="MS-EV-301",
      place="env", cells=("<= 18", "%"),
      expected="match", reason="상한 18% 일치"),
    C(seq=660, ko="고온동작△V (2hr)", en="Voltage drop in high temperature operation (2 h)",
      upper=0.022, unit="V", level="A", remark="45℃ 동작", method="MS-EV-301",
      expected="missing", reason="영문 문서에 2시간 전압 강하 없음"),
    C(seq=670, ko="고온동작△V (24hr)", en="Voltage drop in high temperature operation (24 h)",
      upper=0.032, unit="V", level="A", remark="45℃ 동작", method="MS-EV-301",
      expected="missing", reason="영문 문서에 24시간 전압 강하 없음"),
    C(seq=680, ko="고온동작△T (2hr) 3", en="Thickness increase in high temperature operation (2 h)",
      upper=7, unit="%", level="A", remark="초기 두께 대비", method="MS-EV-301",
      place="env", cells=("<= 0.30", "mm"),
      expected="unknown", reason="기준은 비율(7%), 대상은 절대두께(0.30mm) — 초기 두께 없이는 환산 불가"),
    C(seq=690, ko="고온동작 후  용량(0.2C)", en="Capacity after high temperature operation (0.2C)",
      lower=94, unit="%", level="B", remark="초기 용량 대비", method="MS-EV-301",
      place="env", cells=(">= 94", "%"),
      expected="match", reason="하한 94% 일치"),
    C(seq=700, ko="고온동작 후 용량(1.0C)", en="Capacity after high temperature operation (1.0C)",
      lower=92, unit="%", level="B", remark="초기 용량 대비", method="MS-EV-301",
      place="env", cells=(">= 90", "%"),
      expected="mismatch", mismatch=("lower_limit",),
      reason="기준 하한 92% vs 대상 90%"),
    C(seq=710, ko="Dyne Pen 검사", en="Dyne pen test",
      lower=38, unit="dyn/cm", level="C", remark="파우치 표면 젖음성", method="MS-SF-001",
      place="bullet",
      text="The pouch surface shall show a wetting tension of 38 dyn/cm or higher.",
      expected="match", reason="하한 38 dyn/cm 일치"),

    # 신뢰성특성 / Mass ILT/PRT ---------------------------------------------
    C(seq=720, ko="OCV", en="Open circuit voltage (ILT)",
      lower=3.80, target=3.83, upper=3.86, unit="V", level="B",
      remark="양산 전수 검사", method="MS-PRT-001",
      place="ilt",
      text="Open circuit voltage: 3.80 to 3.86 V (center 3.83 V)",
      expected="match", reason="하한·중심·상한 모두 일치"),
    C(seq=730, ko="△IR", en="IR deviation (ILT)",
      upper=3, unit="mΩ", level="B", remark="Lot 내 편차", method="MS-PRT-001",
      place="ilt",
      text="IR deviation: 3 mOhm or less within a lot",
      expected="match", reason="상한 3mΩ 일치"),
    C(seq=740, ko="Short", en="Internal short (ILT)",
      lower=100, unit="MΩ", level="A", remark="500V 절연 저항", method="MS-PRT-002",
      place="ilt",
      text="Internal short: insulation resistance of 100 MOhm or more at 500 V",
      expected="match", reason="하한 100MΩ·인가전압 500V 일치"),
    C(seq=750, ko="Swelling", en="Swelling (ILT)",
      upper=0.10, unit="mm", level="A", remark="검사 전후 두께차", method="MS-PRT-003",
      place="ilt",
      text="Swelling: 0.15 mm or less",
      expected="mismatch", mismatch=("upper_limit",),
      reason="기준 상한 0.10mm vs 대상 0.15mm"),

    # 신뢰성특성 / 전기적 특성 ----------------------------------------------
    C(seq=760, ko="표준용량", en="Standard capacity", lower=1150, target=1180,
      unit="mAh", level="A", remark="0.2C 방전", method="MS-CP-001",
      place="elec", cells=(">= 1150 (typ. 1180)", "mAh", "MS-CP-001"),
      text=" |  | Typ. | 1,185mAh",   # 2.3.1 연속 행 — 앞 두 칸은 원문대로 비어 있다
      expected="mismatch", mismatch=("target_value",),
      reason="대상 문서가 자기모순 — 6절 표는 typ. 1180, 2.3.1 은 Typ. 1,185mAh. 기준 1180 과 어긋나는 진술이 있음"),
    C(seq=770, ko="출하용량", en="Shipping capacity", lower=340, target=355,
      unit="mAh", level="B", remark="출하 SOC 30% 기준", method="MS-CP-002",
      place="elec", cells=("About 30 % of the rated capacity", "-", "MS-CP-002"),
      expected="unknown", reason="대상이 비율로만 서술하고 mAh 수치를 제시하지 않아 대조 불가"),
    C(seq=780, ko="출하검사용량2", en="Shipping inspection capacity", lower=1140,
      unit="mAh", level="B", remark="샘플 검사", method="MS-CP-003",
      expected="missing", reason="영문 문서에 출하 검사 용량 없음"),
    C(seq=790, ko="정격용량", en="Rated capacity (electrical)", lower=1150,
      unit="mAh", level="A", remark="기본사양 30번과 동일 항목", method="MS-CP-001",
      place="elec", cells=(">= 1150", "mAh", "MS-CP-001"),
      expected="match", reason="1150 일치 (기준에 동명 항목이 둘 — 1:N 케이스)"),
    C(seq=800, ko="GB/T용량 (1.0C충전)", en="GB/T capacity (1.0C charge)", lower=1130,
      unit="mAh", level="B", remark="GB/T 18287 준용", method="GB/T 18287",
      place="elec", cells=(">= 1130", "mAh", "GB/T 18287"),
      expected="match", reason="하한 1130mAh·규격 GB/T 18287 일치"),
    C(seq=810, ko="율별방전 : 0.2C", en="Rate discharge: 0.2C", lower=99, unit="%",
      level="B", remark="표준용량 대비", method="MS-RT-001",
      place="rate", cells=(">= 99 %",),
      expected="match", reason="하한 99% 일치"),
    C(seq=820, ko="율별방전 : 0.5C", en="Rate discharge: 0.5C", lower=97, unit="%",
      level="B", remark="표준용량 대비", method="MS-RT-001",
      expected="missing", reason="영문 문서에 0.5C 율별방전 없음"),
    C(seq=830, ko="율별방전 : 1.0C", en="Rate discharge: 1.0C", lower=95, unit="%",
      level="B", remark="표준용량 대비", method="MS-RT-001",
      place="rate", cells=(">= 93 %",),
      expected="mismatch", mismatch=("lower_limit",),
      reason="기준 하한 95% vs 대상 93%"),
    C(seq=840, ko="온도별방전(3.0V) :  -20℃", en="Temperature discharge at -20 °C",
      lower=55, unit="%", level="B", remark="종지 3.0V", method="MS-RT-002",
      place="rate", cells=(">= 55 %",),
      expected="match", reason="하한 55% 일치"),
    C(seq=850, ko="온도별방전(3.0V) :  -10℃", en="Temperature discharge at -10 °C",
      lower=70, unit="%", level="B", remark="종지 3.0V", method="MS-RT-002",
      place="rate", cells=(">= 70 %",),
      expected="match", reason="하한 70% 일치"),
    C(seq=860, ko="온도별방전(3.0V) :  55℃", en="Temperature discharge at 55 °C",
      lower=96, unit="%", level="B", remark="종지 3.0V", method="MS-RT-002",
      place="rate", cells=(">= 96 %",),
      expected="match", reason="하한 96% 일치"),
    C(seq=870, ko="만충전 상온 28일 방치  (잔존용량)",
      en="Residual capacity after 28 day storage at room temperature (full charge)",
      lower=90, unit="%", level="B", remark="만충전 후 방치", method="MS-ST-001",
      place="para",
      text="After 28 days of storage at room temperature in the fully charged state, "
           "the residual capacity shall be 90 % or more of the initial capacity.",
      expected="match", reason="하한 90%·조건(만충전·상온 28일) 일치"),
    C(seq=880, ko="만충전 상온 28일 방치  (회복용량)",
      en="Recovered capacity after 28 day storage at room temperature (full charge)",
      lower=95, unit="%", level="B", remark="재충전 후 측정", method="MS-ST-001",
      place="para",
      text="After recharging, the recovered capacity shall be 95 % or more of the "
           "initial capacity.",
      expected="match", reason="하한 95% 일치"),
    C(seq=890, ko="출하충전 상온 12달  방치(잔존용량)",
      en="Residual capacity after 12 month storage at room temperature",
      lower=85, unit="%", level="B", remark="출하 SOC 기준", method="MS-ST-002",
      place="store", cells=("12 months, room temperature", ">= 85 %", "-"),
      expected="match", reason="하한 85%·기간 12개월 일치"),
    C(seq=900, ko="출하충전 고온(45℃) 60일 방치 (잔존용량)",
      en="Residual capacity after 60 day storage at 45 °C",
      lower=80, unit="%", level="B", remark="고온 가속", method="MS-ST-003",
      place="store", cells=("60 days, 45 °C", ">= 82 %", "-"),
      expected="mismatch", mismatch=("lower_limit",),
      reason="기준 하한 80% vs 대상 82%"),
    C(seq=910, ko="출하충전상온  1달방치(회복용량)",
      en="Recovered capacity after 1 month storage", lower=97, unit="%", level="C",
      remark="재충전 후 측정", method="MS-ST-002",
      place="store", cells=("1 month, room temperature", "-", ">= 97 %"),
      expected="match", reason="하한 97% 일치"),
    C(seq=920, ko="출하충전상온  2달방치(회복용량)",
      en="Recovered capacity after 2 month storage", lower=96, unit="%", level="C",
      remark="재충전 후 측정", method="MS-ST-002",
      place="store", cells=("2 months, room temperature", "-", ">= 96 %"),
      expected="match", reason="하한 96% 일치"),
    C(seq=930, ko="출하충전상온 3달방치(회복용량)",
      en="Recovered capacity after 3 month storage", lower=95, unit="%", level="C",
      remark="재충전 후 측정", method="MS-ST-002",
      expected="missing", reason="영문 문서에 3개월 회복용량 없음"),
    C(seq=940, ko="출하충전상온  6달방치(회복용량)",
      en="Recovered capacity after 6 month storage", lower=93, unit="%", level="C",
      remark="재충전 후 측정", method="MS-ST-002",
      expected="missing", reason="영문 문서에 6개월 회복용량 없음"),
    C(seq=950, ko="1개월 출충(선박심도)  방치 후 용량측정(GB/T)",
      en="Capacity after 1 month marine transport storage (GB/T)", lower=96, unit="%",
      level="C", remark="선박 수송 조건", method="GB/T 18287",
      expected="missing", reason="영문 문서에 선박 수송 1개월 조건 없음"),
    C(seq=960, ko="2개월 출충(선박심도)  방치 후 용량측정(GB/T)",
      en="Capacity after 2 month marine transport storage (GB/T)", lower=95, unit="%",
      level="C", remark="선박 수송 조건", method="GB/T 18287",
      expected="missing", reason="영문 문서에 선박 수송 2개월 조건 없음"),
    C(seq=970, ko="3개월 출충(선박심도) 방치 후 용량측정(GB/T)",
      en="Capacity after 3 month marine transport storage (GB/T)", lower=94, unit="%",
      level="C", remark="선박 수송 조건", method="GB/T 18287",
      place="store", cells=("3 months, marine transport", ">= 94 %", "-"),
      expected="match", reason="하한 94%·기간 3개월 일치"),
    C(seq=980, ko="6개월 출충(선박심도)  방치 후 용량측정(GB/T)",
      en="Capacity after 6 month marine transport storage (GB/T)", lower=92, unit="%",
      level="C", remark="선박 수송 조건", method="GB/T 18287",
      expected="missing", reason="영문 문서에 선박 수송 6개월 조건 없음"),
    C(seq=990, ko="12개월 출충(선박심도) 방치 후 용량측정(GB/T)",
      en="Capacity after 12 month marine transport storage (GB/T)", lower=89, unit="%",
      level="C", remark="선박 수송 조건", method="GB/T 18287",
      place="store", cells=("12 months, marine transport", ">= 85 %", "-"),
      expected="mismatch", mismatch=("lower_limit",),
      reason="기준 하한 89% vs 대상 85%"),

    # 신뢰성특성 / 고온방치특성 ---------------------------------------------
    C(seq=1000, ko="고온연속충전 두께 3", en="Thickness under continuous charge at high temperature",
      upper=5.20, unit="mm", level="A", remark="60℃ 7일 연속 충전", method="MS-EV-401",
      place="note",
      text="Under continuous charging at high temperature, the cell thickness shall not "
           "exceed 5.20.",
      expected="unknown", reason="대상이 단위를 밝히지 않아 mm 기준과 등가인지 확정 불가"),
]


# ---------------------------------------------------------------------------
# 파생 유틸
# ---------------------------------------------------------------------------

def fmt(v):
    """엑셀 표시용 숫자 문자열 (정수는 소수점 없이)."""
    if v is None:
        return None
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


def spec_string(c):
    """I열 SPEC(중심치 +- 상하한치) 문자열."""
    lo, tg, up = c["lower"], c["target"], c["upper"]
    if lo is not None and up is not None and tg is not None:
        d_lo, d_up = round(tg - lo, 6), round(up - tg, 6)
        if abs(d_lo - d_up) < 1e-9:
            return f"{fmt(tg)} ± {fmt(d_up)}"
        return f"{fmt(tg)} +{fmt(d_up)} / -{fmt(d_lo)}"
    if lo is not None and up is not None:
        return f"{fmt(lo)} ~ {fmt(up)}"
    if lo is not None and tg is not None:
        return f"≥ {fmt(lo)} (typ. {fmt(tg)})"
    if up is not None:
        return f"≤ {fmt(up)}"
    if lo is not None:
        return f"≥ {fmt(lo)}"
    if tg is not None:
        return fmt(tg)
    return None


def new_cases():
    return [c for c in CASES if c["seq"] is not None]


def attributes(c):
    """골든 레코드의 reference.attributes."""
    a = {}
    if c["lower"] is not None:
        a["lower_limit"] = c["lower"]
    if c["target"] is not None:
        a["target_value"] = c["target"]
    if c["upper"] is not None:
        a["upper_limit"] = c["upper"]
    if c["unit"]:
        a["unit"] = c["unit"]
    if c["qual"]:
        a["qualitative_spec"] = c["qual"].replace("\n", " ")
    return a


def target_text(c):
    """정답지의 target_text — 워드 문서에 실제로 들어간 문자열과 같아야 한다."""
    if c["expected"] == "missing":
        return None
    if c["text"] is not None:
        return c["text"]
    # store 표에는 항목명 열이 없다 (조건 문구가 곧 첫 열)
    if c["place"] == "store":
        return " | ".join(c["cells"])
    return " | ".join((c["en"],) + tuple(c["cells"]))


# ---------------------------------------------------------------------------
# 1) 기준 엑셀 값 채우기
# ---------------------------------------------------------------------------

def fill_reference(path=REF_XLSX):
    import openpyxl  # 코어 의존성이 아니라 지연 import 한다 (런타임은 xlwings 를 쓴다)

    wb = openpyxl.load_workbook(path)
    ws = wb[REF_SHEET]

    by_seq = {}
    for r in range(REF_FIRST_NEW_ROW, ws.max_row + 1):
        seq = ws.cell(r, COL["seq"]).value
        if seq is not None:
            by_seq[int(seq)] = r

    filled = 0
    for c in new_cases():
        row = by_seq.get(c["seq"])
        if row is None:
            sys.exit(f"[중단] 순번 {c['seq']} ({c['ko']}) 에 해당하는 행이 없다")
        if ws.cell(row, COL["sub"]).value != c["ko"]:
            sys.exit(f"[중단] {row}행 소분류 불일치: 기대 {c['ko']!r} / "
                     f"실제 {ws.cell(row, COL['sub']).value!r}")
        ws.cell(row, COL["lower"]).value = c["lower"]
        ws.cell(row, COL["target"]).value = c["target"]
        ws.cell(row, COL["upper"]).value = c["upper"]
        ws.cell(row, COL["spec"]).value = spec_string(c)
        ws.cell(row, COL["qual"]).value = c["qual"]
        ws.cell(row, COL["unit"]).value = c["unit"]
        ws.cell(row, COL["level"]).value = c["level"]
        ws.cell(row, COL["remark"]).value = c["remark"]
        ws.cell(row, COL["method"]).value = c["method"]
        filled += 1

    wb.save(path)
    print(f"saved: {path} ({filled} rows filled, rows 3~26 untouched)")


# ---------------------------------------------------------------------------
# 2) 워드 문서 생성
# ---------------------------------------------------------------------------

def _set_cell(cell, text):
    """셀 안의 ``\\n`` 을 <w:br/> 로 넣는다 — 별도 문단으로 쪼개면 표의 행 구조가 흐려진다."""
    lines = text.split("\n")
    run = cell.paragraphs[0].add_run(lines[0])
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def _table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    for ci, h in enumerate(header):
        _set_cell(t.cell(0, ci), h)
    for ri, values in enumerate(rows, start=1):
        for ci, v in enumerate(values):
            _set_cell(t.cell(ri, ci), v)
    return t


def _placed(place):
    return [c for c in CASES if c["place"] == place]


def build_word(path=OUT_DOCX):
    doc = Document()

    for kind, text in LEGACY_PARAGRAPHS:
        if kind == "h1":
            doc.add_heading(text, level=1)
        elif kind == "h2":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)

    # 4구간 충전 온도 — 한 문단 안에서 <w:br/> 로 줄만 나눈다 (c27932e 회귀 케이스)
    p = doc.add_paragraph()
    run = p.add_run(LEGACY_CHARGE_TEMP[0])
    for line in LEGACY_CHARGE_TEMP[1:]:
        run.add_break()
        run.add_text(line)

    doc.add_heading(LEGACY_TABLE_HEADING, level=2)
    _table(doc, LEGACY_TABLE[0], LEGACY_TABLE[1:])

    doc.add_heading(GENERAL_SPEC_HEADING, level=3)
    _table(doc, GENERAL_SPEC_TABLE[0], GENERAL_SPEC_TABLE[1:])

    # 3. 측정/시료 — 서술 문단
    doc.add_heading("3. Measurement and Sampling", level=2)
    for c in _placed("method"):
        doc.add_paragraph(c["text"])

    # 4. 치수/중량 — 5열 표 (앞의 3열 표와 형태를 일부러 다르게 한다)
    doc.add_heading("4. Dimensions and Weight", level=2)
    doc.add_paragraph("All dimensions are measured at 25 °C after the cell has been "
                      "left at rest for 2 hours.")
    _table(doc, ("Item", "Min", "Typ", "Max", "Unit"),
           [(c["en"],) + tuple(c["cells"]) for c in _placed("dim")])

    # 5. 외관/공정 — 불릿
    doc.add_heading("5. Appearance and Workmanship", level=2)
    for c in _placed("bullet"):
        doc.add_paragraph(c["text"], style="List Bullet")

    # 6. 전기적 성능 — 4열 표 (Method 열 추가)
    doc.add_heading("6. Electrical Performance", level=2)
    _table(doc, ("Item", "Specification", "Unit", "Method"),
           [(c["en"],) + tuple(c["cells"]) for c in _placed("elec")])

    # 7. 율별/온도별 방전 — 2열 표
    doc.add_heading("7. Rate and Temperature Discharge", level=2)
    doc.add_paragraph("Retention is expressed as a percentage of the standard capacity.")
    _table(doc, ("Condition", "Retention"),
           [(c["en"],) + tuple(c["cells"]) for c in _placed("rate")])

    # 8. 환경 시험 — 3열 표
    doc.add_heading("8. Environmental Test Results", level=2)
    _table(doc, ("Test item", "Criteria", "Unit"),
           [(c["en"],) + tuple(c["cells"]) for c in _placed("env")])

    # 9. 저장 특성 — 서술 문단 + 3열 표 (항목명 열이 없는 형태)
    doc.add_heading("9. Storage Characteristics", level=2)
    for c in _placed("para"):
        doc.add_paragraph(c["text"])
    _table(doc, ("Storage condition", "Residual capacity", "Recovered capacity"),
           [tuple(c["cells"]) for c in _placed("store")])

    # 10. 양산 검사 — "Term: value" 정의형 문단
    doc.add_heading("10. Mass Production Inspection (ILT/PRT)", level=2)
    doc.add_paragraph("The following items are inspected on every production lot.")
    for c in _placed("ilt"):
        doc.add_paragraph(c["text"])

    # 11. 비고 — 불릿 + 기존 FOOTER
    doc.add_heading("11. Notes", level=2)
    for c in _placed("note"):
        doc.add_paragraph(c["text"], style="List Bullet")
    doc.add_paragraph(FOOTER)

    doc.save(path)
    print(f"saved: {path}")


# ---------------------------------------------------------------------------
# 3) 정답지 생성
# ---------------------------------------------------------------------------

def write_golden(path=OUT_GOLDEN, ref_rows=None):
    """ref_rows: {순번: 엑셀 행}. 없으면 엑셀에서 읽는다."""
    if ref_rows is None:
        import openpyxl
        ws = openpyxl.load_workbook(REF_XLSX)[REF_SHEET]
        ref_rows = {}
        for r in range(REF_FIRST_NEW_ROW, ws.max_row + 1):
            seq = ws.cell(r, COL["seq"]).value
            if seq is not None:
                ref_rows[int(seq)] = r

    records = []
    for i, c in enumerate(CASES, start=1):
        row = c["row"] if c["row"] is not None else ref_rows[c["seq"]]
        records.append({
            "id": f"g-en-{i:03d}",
            "entity_name": c["ko"],
            "target_doc": "spec_en.docx",
            "reference": {
                "doc": "자표준문서.xlsx",
                "row": row,
                "cell_range": f"E{row}:N{row}",
                "attributes": attributes(c),
            },
            "target_text": target_text(c),
            "expected": c["expected"],
            "mismatch_attributes": list(c["mismatch"]),
            "reason": c["reason"],
        })

    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        for rec in records:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")

    from collections import Counter
    dist = Counter(r["expected"] for r in records)
    print(f"saved: {path} ({len(records)} records, "
          + " / ".join(f"{k} {dist[k]}" for k in ("match", "mismatch", "unknown", "missing"))
          + ")")


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--skip-reference", action="store_true",
                    help="기준 엑셀(samples/자표준문서.xlsx)은 건드리지 않는다")
    args = ap.parse_args()

    if not args.skip_reference:
        fill_reference()
    build_word()
    write_golden()


if __name__ == "__main__":
    main()
