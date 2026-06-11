"""Word Raw Extractor 테스트.

python-docx 로 임시 .docx 를 만든 뒤 :func:`extract_word_raw` 가 문단/표를
문서 순서대로 physical_raw 로 만드는지 검증한다. 기획의 '충전환경온도' 설명 문단 +
규격표 흐름을 재현한다.

python-docx 미설치 환경에서는 전체 모듈을 skip 한다.
"""

from __future__ import annotations

import pytest

docx = pytest.importorskip("docx")

from contentcompare.raw import extract_raw, raw_to_json
from contentcompare.raw.word_raw import extract_word_raw


@pytest.fixture()
def standard_docx(tmp_path):
    """heading → 설명 문단 → 규격표 순서의 Word 문서."""
    from docx import Document

    document = Document()
    document.add_heading("기본사양", level=1)  # 첫 블록: heading 문단
    document.add_paragraph(
        "충전환경온도는 -5℃에서 55℃ 범위로 관리하며, 중심치는 25℃로 한다."
    )

    table = document.add_table(rows=2, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "항목", "규격", "단위"
    data = table.rows[1].cells
    data[0].text, data[1].text, data[2].text = "충전환경온도", "-5~55", "℃"

    document.add_paragraph("")  # 빈 문단 — 추출에서 제외되어야 한다

    path = tmp_path / "standard_description.docx"
    document.save(path)
    return str(path)


def test_doc_root_shape(standard_docx):
    doc = extract_word_raw(standard_docx).to_dict()
    assert doc["doc_type"] == "word"
    assert doc["file_name"] == "standard_description.docx"
    assert isinstance(doc["blocks"], list)


def test_blocks_in_document_order(standard_docx):
    blocks = extract_word_raw(standard_docx).to_dict()["blocks"]
    # heading(문단) → 설명(문단) → 표. 빈 문단은 제외 → 정확히 3개.
    assert len(blocks) == 3
    assert [b["type"] for b in blocks] == ["paragraph", "paragraph", "table"]
    # order 가 1,2,3 으로 연속.
    assert [b["order"] for b in blocks] == [1, 2, 3]
    assert [b["block_id"] for b in blocks] == ["w_b001", "w_b002", "w_b003"]


def test_heading_paragraph_style(standard_docx):
    blocks = extract_word_raw(standard_docx).to_dict()["blocks"]
    head = blocks[0]
    assert head["type"] == "paragraph"
    assert head["text"] == "기본사양"
    assert "Heading" in head["style"]["style_name"]


def test_description_paragraph_text(standard_docx):
    blocks = extract_word_raw(standard_docx).to_dict()["blocks"]
    desc = blocks[1]
    assert desc["type"] == "paragraph"
    assert "충전환경온도" in desc["text"]
    assert "-5℃에서 55℃" in desc["text"]


def test_table_rows_extracted(standard_docx):
    blocks = extract_word_raw(standard_docx).to_dict()["blocks"]
    table = blocks[2]
    assert table["type"] == "table"
    assert table["rows"][0] == ["항목", "규격", "단위"]
    assert table["rows"][1] == ["충전환경온도", "-5~55", "℃"]


def test_empty_paragraph_skipped(standard_docx):
    blocks = extract_word_raw(standard_docx).to_dict()["blocks"]
    assert all((b.get("text") or "").strip() for b in blocks if b["type"] == "paragraph")


def test_no_interpretation_only_physical(standard_docx):
    text = raw_to_json(extract_word_raw(standard_docx))
    assert "entity" not in text
    assert "lower_limit" not in text
    assert "block_id" in text


def test_dispatcher_routes_docx(standard_docx):
    doc = extract_raw(standard_docx).to_dict()
    assert doc["doc_type"] == "word"
